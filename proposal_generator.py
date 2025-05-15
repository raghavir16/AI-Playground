import os
import autogen
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Flag to use mock LLM (for testing without API key)
USE_MOCK_LLM = False  # Set to False for real OpenAI, True for testing

# Configuration for the LLM
if USE_MOCK_LLM:
    # Mock configuration for testing
    config_list = [{
        "model": "gpt-3.5-turbo",
        "api_key": "mock_key",
    }]
    
    # LLM config for testing
    llm_config = {
        "config_list": config_list,
        "temperature": 0,
    }
else:
    # Real OpenAI configuration
    config_list = [
        {
            "model": "gpt-4.1",
            "api_key": os.environ.get("OPENAI_API_KEY"),
        }
    ]
    
    # Define the LLM configuration that we'll use for all agents
    llm_config = {
        "config_list": config_list,
        "seed": 42,  # for reproducibility
        "temperature": 0.7,  # slightly creative but not too random
    }

# Function to validate the API key
def validate_api_key():
    if USE_MOCK_LLM:
        return True  # Skip validation for mock LLM
    
    if not os.environ.get("OPENAI_API_KEY"):
        raise ValueError(
            "OPENAI_API_KEY environment variable not set. "
            "Please set it with your OpenAI API key."
        )

# Function to create a Word document from proposal content
def create_proposal_document(content, template_path=None, output_path="proposal.docx"):
    """
    Creates a Word document using the provided content and template
    
    Args:
        content (dict): Dictionary containing content for each section
        template_path (str): Path to the Word template
        output_path (str): Path where the output document will be saved
    
    Returns:
        str: Path to the generated document
    """
    try:
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            logger.warning("Template not found, using default Word document")
        
        # Title Page
        doc.add_heading(content.get("title", "IT Project Proposal"), 0)
        doc.add_paragraph(content.get("subtitle", ""))
        doc.add_page_break()
        
        # Table of Contents placeholder
        doc.add_heading("Contents", 1)
        doc.add_paragraph("Contents will be generated automatically")
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(content.get("executive_summary", ""))
        
        # Customer Requirements Table
        doc.add_heading("Customer Requirements", 1)
        requirements = content.get("requirements", [])
        if requirements:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Requirement'
            header_cells[1].text = 'Description'
            
            for req in requirements:
                row_cells = table.add_row().cells
                row_cells[0].text = req.get("requirement", "")
                row_cells[1].text = req.get("description", "")
        
        # Scope Statements
        doc.add_heading("Scope", 1)
        doc.add_heading("In Scope", 2)
        for item in content.get("in_scope", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Out of Scope", 2)
        for item in content.get("out_scope", []):
            doc.add_paragraph(item, style='List Bullet')
        # Mandatory out of scope statement
        doc.add_paragraph("Anything not specifically mentioned in the In-Scope section is considered Out of Scope.", style='List Bullet')
        
        # Solution Summary
        doc.add_heading("Solution Summary", 1)
        doc.add_paragraph(content.get("solution_summary", ""))
        
        # Deliverables
        doc.add_heading("Deliverables", 1)
        doc.add_heading("Standard Deliverables", 2)
        for item in content.get("standard_deliverables", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Project-Specific Deliverables", 2)
        for item in content.get("project_specific_deliverables", []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Costs
        doc.add_heading("Costs", 1)
        doc.add_heading("Resource Costs", 2)
        resources = content.get("resources", [])
        if resources:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Activity'
            header_cells[1].text = 'Role Type'
            header_cells[2].text = 'Quantity (days)'
            header_cells[3].text = 'Unit Cost'
            header_cells[4].text = 'Total Cost'
            
            for res in resources:
                row_cells = table.add_row().cells
                row_cells[0].text = res.get("activity", "")
                row_cells[1].text = res.get("role_type", "")
                row_cells[2].text = str(res.get("quantity", ""))
                row_cells[3].text = res.get("unit_cost", "")
                row_cells[4].text = res.get("total_cost", "")
        
        doc.add_heading("Licensing", 2)
        doc.add_paragraph(content.get("licensing", "No additional licensing required"))
        
        # RAID Section
        doc.add_heading("RAID", 1)
        doc.add_heading("Risks", 2)
        for item in content.get("risks", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Assumptions", 2)
        for item in content.get("assumptions", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Issues", 2)
        for item in content.get("issues", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Dependencies", 2)
        for item in content.get("dependencies", []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Tasks and Effort
        doc.add_heading("Tasks and Effort Estimates", 1)
        tasks = content.get("tasks", [])
        if tasks:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Task'
            header_cells[1].text = 'Effort (days)'
            
            for task in tasks:
                row_cells = table.add_row().cells
                row_cells[0].text = task.get("task", "")
                row_cells[1].text = str(task.get("effort", ""))
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Proposal document created successfully at {output_path}")
        print(f"\n\n========================")
        print(f"SUCCESS: Proposal document created successfully at {output_path}")
        print(f"========================\n")
        return output_path
    except Exception as e:
        error_msg = f"Error creating document: {str(e)}"
        logger.error(error_msg)
        print(f"\n\n========================")
        print(f"ERROR: {error_msg}")
        print(f"========================\n")
        return None

# Define the agents

def create_agents(config_list, template_path=None):
    """
    Creates and returns the agents needed for proposal generation
    
    Args:
        config_list: Configuration for the LLM
        template_path: Path to the Word template
    
    Returns:
        tuple: The created agents
    """
    # Prepare workdir
    if not os.path.exists("workdir"):
        os.makedirs("workdir")
        
    # Copy the create_proposal_document function to workdir for the Document_Assembler to use
    with open("workdir/document_utils.py", "w") as f:
        f.write("""
import os
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_proposal_document(content, template_path=None, output_path="proposal.docx"):
    \"""
    Creates a Word document using the provided content and template
    
    Args:
        content (dict): Dictionary containing content for each section
        template_path (str): Path to the template
        output_path (str): Path where the output document will be saved
    
    Returns:
        str: Path to the generated document
    \"""
    try:
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            print("No template found, using default Word document")
        
        # Title Page
        doc.add_heading(content.get("title", "IT Project Proposal"), 0)
        doc.add_paragraph(content.get("subtitle", ""))
        doc.add_page_break()
        
        # Table of Contents placeholder
        doc.add_heading("Contents", 1)
        doc.add_paragraph("Contents will be generated automatically")
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(content.get("executive_summary", ""))
        
        # Customer Requirements Table
        doc.add_heading("Customer Requirements", 1)
        requirements = content.get("requirements", [])
        if requirements:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Requirement'
            header_cells[1].text = 'Description'
            
            for req in requirements:
                row_cells = table.add_row().cells
                row_cells[0].text = req.get("requirement", "")
                row_cells[1].text = req.get("description", "")
        
        # Scope Statements
        doc.add_heading("Scope", 1)
        doc.add_heading("In Scope", 2)
        for item in content.get("in_scope", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Out of Scope", 2)
        for item in content.get("out_scope", []):
            doc.add_paragraph(item, style='List Bullet')
        # Mandatory out of scope statement
        doc.add_paragraph("Anything not specifically mentioned in the In-Scope section is considered Out of Scope.", style='List Bullet')
        
        # Solution Summary
        doc.add_heading("Solution Summary", 1)
        doc.add_paragraph(content.get("solution_summary", ""))
        
        # Deliverables
        doc.add_heading("Deliverables", 1)
        doc.add_heading("Standard Deliverables", 2)
        for item in content.get("standard_deliverables", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Project-Specific Deliverables", 2)
        for item in content.get("project_specific_deliverables", []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Costs
        doc.add_heading("Costs", 1)
        doc.add_heading("Resource Costs", 2)
        resources = content.get("resources", [])
        if resources:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Activity'
            header_cells[1].text = 'Role Type'
            header_cells[2].text = 'Quantity (days)'
            header_cells[3].text = 'Unit Cost'
            header_cells[4].text = 'Total Cost'
            
            for res in resources:
                row_cells = table.add_row().cells
                row_cells[0].text = res.get("activity", "")
                row_cells[1].text = res.get("role_type", "")
                row_cells[2].text = str(res.get("quantity", ""))
                row_cells[3].text = res.get("unit_cost", "")
                row_cells[4].text = res.get("total_cost", "")
        
        doc.add_heading("Licensing", 2)
        doc.add_paragraph(content.get("licensing", "No additional licensing required"))
        
        # RAID Section
        doc.add_heading("RAID", 1)
        doc.add_heading("Risks", 2)
        for item in content.get("risks", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Assumptions", 2)
        for item in content.get("assumptions", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Issues", 2)
        for item in content.get("issues", []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Dependencies", 2)
        for item in content.get("dependencies", []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Tasks and Effort
        doc.add_heading("Tasks and Effort Estimates", 1)
        tasks = content.get("tasks", [])
        if tasks:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Task'
            header_cells[1].text = 'Effort (days)'
            
            for task in tasks:
                row_cells = table.add_row().cells
                row_cells[0].text = task.get("task", "")
                row_cells[1].text = str(task.get("effort", ""))
        
        # Save the document
        doc.save(output_path)
        print(f"\\n\\n========================")
        print(f"SUCCESS: Proposal document created successfully at {{output_path}}")
        print(f"========================\\n")
        return output_path
    except Exception as e:
        error_msg = f"Error creating document: {{str(e)}}"
        print(f"\\n\\n========================")
        print(f"ERROR: {{error_msg}}")
        print(f"========================\\n")
        return None
""")
    
    # User Proxy Agent
    user_proxy = autogen.UserProxyAgent(
        name="User_Proxy",
        system_message="You are a human user interacting with the AI proposal system. You'll provide the initial project request and review the final output. When the proposal is complete, ask for the Microsoft Word document to be generated if it hasn't been already. The final deliverable must be a Word document file, not just text content.",
        human_input_mode="ALWAYS",
        code_execution_config={"last_n_messages": 3, "work_dir": "workdir"},
    )

    # Architect Agent
    architect = autogen.AssistantAgent(
        name="Architect",
        system_message="You are an IT Architect working on a project proposal. Your role is to provide the initial project overview and answer questions from the proposal team based on your technical expertise. You'll also review drafts and provide feedback on technical accuracy and completeness.",
        llm_config=llm_config,
    )

    # Proposal Manager Agent
    proposal_manager = autogen.AssistantAgent(
        name="Proposal_Manager",
        system_message="""You are the Proposal Manager responsible for:
1. Receiving the initial prompt from the Architect.
2. Coordinating the overall workflow.
3. Delegating tasks to specialized agents.
4. Managing the conversation flow.
5. Synthesizing information and requesting clarification when needed.
6. Presenting the final document to the Architect.
7. Routing feedback to appropriate agents for revision.

Always start by understanding the project overview, then coordinate the other agents to complete their tasks.

When all information has been collected, explicitly instruct the Document_Assembler to call the create_proposal_document function to generate a Microsoft Word document with the compiled content. Ensure that an actual document file is created, not just a text description of the document content.""",
        llm_config=llm_config,
    )

    # Requirements Analyst Agent
    requirements_analyst = autogen.AssistantAgent(
        name="Requirements_Analyst",
        system_message="""You are the Requirements Analyst responsible for:
1. Analyzing the initial prompt and identifying missing information.
2. Formulating clear questions to gather missing details.
3. Formatting the gathered requirements into a structured table.
4. Consulting the Design Standards document if provided.

Focus on identifying critical requirements including Hardware, Software, Volumes, Licensing, and Prerequisites.""",
        llm_config=llm_config,
    )

    # Solution Designer Agent
    solution_designer = autogen.AssistantAgent(
        name="Solution_Designer",
        system_message="""You are the Solution Designer responsible for:
1. Drafting the Executive Summary based on the project overview.
2. Defining the Scope Statements (In Scope / Out of Scope).
3. Writing the Solution Summary with high-level project overview and activities.
4. Identifying relevant diagrams (if possible).
5. Defining the Deliverables section.
6. Drafting the Tasks and Effort Estimates.

Be specific and focus on technical details while ensuring all information aligns with the project requirements.""",
        llm_config=llm_config,
    )

    # Cost Estimator Agent
    cost_estimator = autogen.AssistantAgent(
        name="Cost_Estimator",
        system_message="""You are the Cost Estimator responsible for:
1. Drafting the Costs section.
2. Generating the resource table with activities, roles, and quantities.
3. Generating the section for required licenses.

Provide realistic estimates based on the solution design and deliverables.""",
        llm_config=llm_config,
    )

    # Risk Assessor Agent
    risk_assessor = autogen.AssistantAgent(
        name="Risk_Assessor",
        system_message="""You are the Risk Assessor responsible for:
1. Drafting the RAID (Risks, Assumptions, Issues, Dependencies) section.
2. Identifying common IT project risks and assumptions.
3. Tailoring them to the specific project context.

Focus on realistic risks that could impact project delivery and success.""",
        llm_config=llm_config,
    )

    # Document Assembler Agent
    document_assembler = autogen.AssistantAgent(
        name="Document_Assembler",
        system_message=f"""You are the Document Assembler responsible for:
1. Generating the Microsoft Word document using python-docx.
2. Assembling content from other agents into the document.
3. Formatting the document according to standards.
4. Handling revisions by regenerating the document with updated content.

You have access to {'a Littlefish branded template' if template_path else 'a standard Word document format'} for creating the proposal document.

CRITICAL INSTRUCTION: You MUST directly EXECUTE the code to create the document, not just show or discuss it. 
Follow these exact steps:
1. Collect all content from the conversation
2. Format it into a clean content dictionary
3. EXECUTE the create_proposal_document function 
4. Confirm document creation to the user

Execute this exact code block structure (with your collected content) when ready:

```python
# This is self-executing code that will run when this message is sent
from document_utils import create_proposal_document

content = {{
    "title": "IT Project Proposal - [YOUR COLLECTED PROJECT NAME]",
    "subtitle": "Prepared for [CLIENT NAME]",
    "executive_summary": "[YOUR COLLECTED EXECUTIVE SUMMARY]",
    "requirements": [
        {{"requirement": "[REQUIREMENT 1]", "description": "[DESCRIPTION 1]"}},
        # Add all requirements collected from the conversation
    ],
    "in_scope": [
        "[ITEM 1]",
        "[ITEM 2]",
        # Add all in-scope items
    ],
    "out_scope": [
        "[ITEM 1]",
        "[ITEM 2]",
        # Add all out-of-scope items
    ],
    "solution_summary": "[YOUR COLLECTED SOLUTION SUMMARY]",
    "standard_deliverables": [
        "[DELIVERABLE 1]",
        "[DELIVERABLE 2]",
        # Add all standard deliverables
    ],
    "project_specific_deliverables": [
        "[SPECIFIC DELIVERABLE 1]",
        "[SPECIFIC DELIVERABLE 2]",
        # Add all project-specific deliverables
    ],
    "resources": [
        {{"activity": "[ACTIVITY 1]", "role_type": "[ROLE 1]", "quantity": [NUMBER], "unit_cost": "[COST]", "total_cost": "[TOTAL]"}},
        # Add all resources
    ],
    "licensing": "[YOUR COLLECTED LICENSING INFO]",
    "risks": [
        "[RISK 1]",
        "[RISK 2]",
        # Add all risks
    ],
    "assumptions": [
        "[ASSUMPTION 1]",
        "[ASSUMPTION 2]",
        # Add all assumptions
    ],
    "issues": [
        "[ISSUE 1]",
        "[ISSUE 2]",
        # Add all issues
    ],
    "dependencies": [
        "[DEPENDENCY 1]",
        "[DEPENDENCY 2]",
        # Add all dependencies
    ],
    "tasks": [
        {{"task": "[TASK 1]", "effort": [NUMBER]}},
        {{"task": "[TASK 2]", "effort": [NUMBER]}},
        # Add all tasks with effort
    ]
}}

# DO NOT modify this line - it will create the document
output_path = create_proposal_document(content)

# DO NOT modify this line - it will report success
print(f"Document created at: {{output_path}}")
```

This code WILL execute when you send this message - no further action is needed. The document will be saved as "proposal.docx" in the working directory. After execution, report to the user that the document has been created and where it can be found.""",
        llm_config=llm_config,
        code_execution_config={
            "last_n_messages": 3, 
            "work_dir": "workdir", 
            "auto_execute": True  # Automatically execute code without asking
        },
    )
    
    # Force the Document_Assembler to execute code without human input
    document_assembler.human_input_mode = "NEVER"
    
    return user_proxy, architect, proposal_manager, requirements_analyst, solution_designer, cost_estimator, risk_assessor, document_assembler

def create_group_chat(agents):
    """Creates a group chat with all agents"""
    user_proxy, architect, proposal_manager, requirements_analyst, solution_designer, cost_estimator, risk_assessor, document_assembler = agents
    
    # Create the group chat
    groupchat = autogen.GroupChat(
        agents=[user_proxy, architect, proposal_manager, requirements_analyst, solution_designer, cost_estimator, risk_assessor, document_assembler],
        messages=[],
        max_round=50,
    )
    
    # Create the group chat manager
    manager = autogen.GroupChatManager(
        groupchat=groupchat,
        llm_config=llm_config,
    )
    
    return manager

def main():
    """Main function to run the proposal generation system"""
    try:
        print("Starting proposal generator...")
        
        if not USE_MOCK_LLM:
            # Ask for API key if not set
            api_key = os.environ.get("OPENAI_API_KEY")
            if not api_key:
                api_key = input("Please enter your OpenAI API key: ").strip()
                os.environ["OPENAI_API_KEY"] = api_key
                
            # Validate API key
            print("Validating API key...")
            validate_api_key()
            
            # Update the config list with the API key
            config_list[0]["api_key"] = api_key
            
        print("Using " + ("mock LLM" if USE_MOCK_LLM else "OpenAI API") + " for chat completions")
        
        # Ask for template path
        template_path = input("Enter the path to the Littlefish Word template (press Enter to use default): ").strip()
        if not template_path:
            template_path = "littlefish_template.docx"
            
        if not os.path.exists(template_path):
            logger.warning(f"Template file not found at {template_path}, using default Word format")
            template_path = None
        
        print("Configuring LLM...")
        
        # Create the agents
        print("Creating agents...")
        agents = create_agents(config_list, template_path)
        
        # Create the group chat
        print("Setting up group chat...")
        manager = create_group_chat(agents)
        
        # Start the conversation
        user_proxy = agents[0]
        
        # Starting message
        print("Initiating chat - this may take a moment...")
        user_proxy.initiate_chat(
            manager,
            message="""
            I need to generate an IT project proposal. I would like to describe my project needs so you can help create a comprehensive proposal document.
            
            Please ask me about:
            - Project description and overview
            - Project budget and financial constraints
            - Ideal timescales and deadlines
            - Technical requirements and specifications
            - Stakeholders and target audience
            - Any existing systems or infrastructure
            - Specific business goals and outcomes
            
            Once you have this information, please create a comprehensive proposal document.
            """
        )
        
       
        output_path = create_proposal_document(content)
        
    except Exception as e:
        logger.error(f"Error: {e}")
        print(f"An error occurred: {e}")
        # Print the full traceback for better debugging
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Error in main program: {e}")
        import traceback
        traceback.print_exc() 
