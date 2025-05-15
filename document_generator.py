from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
import os
from .template_manager import TemplateManager

class ProposalDocumentGenerator:
    def __init__(self, template_path: Optional[str] = None):
        if template_path and os.path.exists(template_path):
            self.template_doc = Document(template_path)
            self.document = Document()
            TemplateManager.apply_template_styles(self.document, self.template_doc)
        else:
            self.document = Document()
            self._setup_default_styles()
    
    def _setup_default_styles(self):
        """Setup default document styles if no template is provided"""
        # Title style
        style = self.document.styles['Title']
        style.font.size = Pt(24)
        style.font.name = 'Arial'
        
        # Heading styles
        for level in range(1, 4):
            style = self.document.styles[f'Heading {level}']
            style.font.name = 'Arial'
            style.font.size = Pt(16 - (level * 2))
    
    def add_title_page(self, title: str, customer: str, date: str):
        """Add the title page"""
        self.document.add_paragraph(title, style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_paragraph(f"Prepared for: {customer}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_paragraph(f"Date: {date}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents"""
        self.document.add_heading('Table of Contents', level=1)
        # Note: Table of contents needs to be updated manually in Word
        self.document.add_paragraph('Right-click and select "Update Field" to update this table of contents')
        self.document.add_page_break()
    
    def add_executive_summary(self, summary: str):
        """Add executive summary section"""
        self.document.add_heading('Executive Summary', level=1)
        self.document.add_paragraph(summary)
        self.document.add_page_break()
    
    def add_requirements_table(self, requirements: List[Dict[str, str]]):
        """Add customer requirements table"""
        self.document.add_heading('Customer Requirements', level=1)
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Requirement ID'
        header_cells[1].text = 'Description'
        header_cells[2].text = 'Priority'
        
        # Add requirement rows
        for req in requirements:
            row_cells = table.add_row().cells
            row_cells[0].text = req['id']
            row_cells[1].text = req['description']
            row_cells[2].text = req['priority']
        
        self.document.add_page_break()
    
    def add_scope_sections(self, in_scope: List[str], out_scope: List[str]):
        """Add scope sections"""
        self.document.add_heading('Project Scope', level=1)
        
        self.document.add_heading('In Scope', level=2)
        for item in in_scope:
            self.document.add_paragraph(item, style='List Bullet')
            
        self.document.add_heading('Out of Scope', level=2)
        self.document.add_paragraph("Anything not specifically mentioned in scope is considered out of scope, including:")
        for item in out_scope:
            self.document.add_paragraph(item, style='List Bullet')
        
        self.document.add_page_break()
    
    def add_solution_summary(self, summary: str, diagram_path: Optional[str] = None):
        """Add solution summary section"""
        self.document.add_heading('Solution Summary', level=1)
        self.document.add_paragraph(summary)
        
        if diagram_path and os.path.exists(diagram_path):
            self.document.add_picture(diagram_path, width=Inches(6))
            
        self.document.add_page_break()
    
    def add_deliverables(self, deliverables: List[Dict[str, str]]):
        """Add deliverables section"""
        self.document.add_heading('Deliverables', level=1)
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Deliverable'
        header_cells[1].text = 'Description'
        
        for deliverable in deliverables:
            row_cells = table.add_row().cells
            row_cells[0].text = deliverable['name']
            row_cells[1].text = deliverable['description']
            
        self.document.add_page_break()
    
    def add_costs_section(self, resources: List[Dict[str, str]], licenses: Optional[List[Dict[str, str]]] = None):
        """Add costs section"""
        self.document.add_heading('Costs and Resources', level=1)
        
        # Resource costs table
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Activity'
        header_cells[1].text = 'Role'
        header_cells[2].text = 'Type'
        header_cells[3].text = 'Days'
        header_cells[4].text = 'Cost'
        
        for resource in resources:
            row_cells = table.add_row().cells
            row_cells[0].text = resource['activity']
            row_cells[1].text = resource['role']
            row_cells[2].text = resource['type']
            row_cells[3].text = str(resource['days'])
            row_cells[4].text = resource.get('cost', '')
            
        if licenses:
            self.document.add_heading('Required Licenses', level=2)
            license_table = self.document.add_table(rows=1, cols=3)
            license_table.style = 'Table Grid'
            
            header_cells = license_table.rows[0].cells
            header_cells[0].text = 'License'
            header_cells[1].text = 'Quantity'
            header_cells[2].text = 'Cost'
            
            for license in licenses:
                row_cells = license_table.add_row().cells
                row_cells[0].text = license['name']
                row_cells[1].text = str(license['quantity'])
                row_cells[2].text = license.get('cost', '')
                
        self.document.add_page_break()
    
    def add_raid_section(self, risks: List[str], assumptions: List[str], 
                        issues: List[str], dependencies: List[str]):
        """Add RAID analysis section"""
        self.document.add_heading('RAID Analysis', level=1)
        
        self.document.add_heading('Risks', level=2)
        for risk in risks:
            self.document.add_paragraph(risk, style='List Bullet')
            
        self.document.add_heading('Assumptions', level=2)
        for assumption in assumptions:
            self.document.add_paragraph(assumption, style='List Bullet')
            
        self.document.add_heading('Issues', level=2)
        for issue in issues:
            self.document.add_paragraph(issue, style='List Bullet')
            
        self.document.add_heading('Dependencies', level=2)
        for dependency in dependencies:
            self.document.add_paragraph(dependency, style='List Bullet')
            
        self.document.add_page_break()
    
    def add_effort_breakdown(self, tasks: List[Dict[str, str]]):
        """Add effort breakdown section"""
        self.document.add_heading('Effort Breakdown', level=1)
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Task'
        header_cells[1].text = 'Description'
        header_cells[2].text = 'Role'
        header_cells[3].text = 'Effort (days)'
        
        for task in tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = task['name']
            row_cells[1].text = task['description']
            row_cells[2].text = task['role']
            row_cells[3].text = str(task['effort'])
    
    def save(self, filename: str):
        """Save the document"""
        self.document.save(filename) 