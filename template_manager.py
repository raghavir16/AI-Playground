from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class TemplateManager:
    @staticmethod
    def create_littlefish_template(output_path: str):
        """Create a Littlefish branded template with predefined styles"""
        doc = Document()
        
        # Set up page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Define brand colors
        LITTLEFISH_BLUE = RGBColor(0, 120, 200)  # Example brand color
        
        # Title Style
        style = doc.styles['Title']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(24)
        font.bold = True
        font.color.rgb = LITTLEFISH_BLUE
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_after = Pt(24)
        
        # Heading Styles
        for level in range(1, 4):
            style = doc.styles[f'Heading {level}']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(16 - (level * 2))
            font.bold = True
            font.color.rgb = LITTLEFISH_BLUE
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        
        # Normal Text Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        style.paragraph_format.space_after = Pt(10)
        
        # Table Style
        table_style = doc.styles.add_style('Littlefish Table', WD_STYLE_TYPE.TABLE)
        table_style._element.set(qn('w:default'), '1')
        
        # List Bullet Style
        list_style = doc.styles['List Bullet']
        font = list_style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Add sample content to demonstrate styles
        doc.add_paragraph('Project Proposal', style='Title')
        doc.add_paragraph('TEMPLATE DOCUMENT - DO NOT MODIFY', style='Subtitle')
        
        # Add headers for each required section
        sections = [
            'Executive Summary',
            'Customer Requirements',
            'Project Scope',
            'Solution Summary',
            'Deliverables',
            'Costs and Resources',
            'RAID Analysis',
            'Effort Breakdown'
        ]
        
        for section in sections:
            doc.add_paragraph(section, style='Heading 1')
            doc.add_paragraph('This section will be populated by the proposal generator.', style='Normal')
            doc.add_paragraph()
        
        # Save the template
        doc.save(output_path)
    
    @staticmethod
    def apply_template_styles(doc: Document, template_doc: Document):
        """Apply styles from template to the target document"""
        # Copy styles from template
        for style in template_doc.styles:
            if style.name not in doc.styles:
                try:
                    doc.styles.add_style(
                        style.name,
                        style.type,
                        base_style=style
                    )
                except ValueError:
                    # Style already exists
                    pass
        
        # Apply table styles
        for table in doc.tables:
            table.style = 'Littlefish Table'
        
        return doc 